from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('AI Startup Idea Validator', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a horizontal line effect (using a table)
doc.add_paragraph()

# Project Overview Section
doc.add_heading('Project Overview', level=1)
overview = doc.add_paragraph(
    'The AI Startup Idea Validator is a full-stack web application that allows users to submit startup ideas '
    'and receive comprehensive AI-powered analysis. The application provides intelligent insights on market demand, '
    'competitors, revenue models, SWOT analysis, and strategic suggestions for entrepreneurs.'
)

doc.add_paragraph()

# System Requirements Section
doc.add_heading('System Requirements', level=1)

doc.add_heading('Server-Side Requirements:', level=2)
server_reqs = [
    'Node.js v18 or higher',
    'npm (Node Package Manager)',
    'Express.js framework',
    'CORS support for cross-origin requests',
    'Minimum 512MB RAM',
    'Port 3000 available for the server'
]
for req in server_reqs:
    doc.add_paragraph(req, style='List Bullet')

doc.add_paragraph()

doc.add_heading('Client-Side Requirements:', level=2)
client_reqs = [
    'Modern web browser (Chrome, Firefox, Safari, Edge)',
    'JavaScript enabled',
    'HTML5 and CSS3 support',
    'Minimum screen resolution: 768x1024 pixels',
    'Internet connection (localhost access)'
]
for req in client_reqs:
    doc.add_paragraph(req, style='List Bullet')

doc.add_paragraph()

# Functional Requirements Section
doc.add_heading('Functional Requirements', level=1)

doc.add_heading('Core Features:', level=2)
features = [
    'User Input Form: Accept startup idea title, description, and target audience',
    'AI Analysis Engine: Generate comprehensive analysis using mock AI (no external API)',
    'Market Demand Analysis: Evaluate market size and growth potential',
    'Competitor Analysis: Identify and analyze competitive landscape',
    'Revenue Model Suggestions: Provide pricing and monetization strategies',
    'SWOT Analysis: Strengths, Weaknesses, Opportunities, and Threats assessment',
    'Strategic Recommendations: Actionable suggestions for startup launch',
    'Input Validation: Ensure users provide meaningful, valid data',
    'Loading Indicator: Visual feedback during analysis processing',
    'Clear Results Button: Allow users to reset and analyze new ideas',
    'Error Handling: Graceful error messages for invalid inputs'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph()

# Non-Functional Requirements Section
doc.add_heading('Non-Functional Requirements', level=1)

nfr_table = doc.add_table(rows=6, cols=2)
nfr_table.style = 'Light Grid Accent 1'

# Header row
header_cells = nfr_table.rows[0].cells
header_cells[0].text = 'Requirement'
header_cells[1].text = 'Details'

# Data rows
requirements_data = [
    ('Performance', 'Response time < 2 seconds for analysis generation'),
    ('Availability', '99% uptime during operational hours'),
    ('Security', 'Input sanitization and validation on server-side'),
    ('Scalability', 'Support for 100+ concurrent users'),
    ('Usability', 'Intuitive UI with responsive design for all devices')
]

for i, (requirement, details) in enumerate(requirements_data, 1):
    row_cells = nfr_table.rows[i].cells
    row_cells[0].text = requirement
    row_cells[1].text = details

doc.add_paragraph()

# Technology Stack Section
doc.add_heading('Technology Stack', level=1)

tech_table = doc.add_table(rows=6, cols=2)
tech_table.style = 'Light Grid Accent 1'

# Header row
tech_header = tech_table.rows[0].cells
tech_header[0].text = 'Component'
tech_header[1].text = 'Technology'

# Data rows
tech_data = [
    ('Backend', 'Node.js with Express.js'),
    ('Frontend', 'HTML5, CSS3, Vanilla JavaScript'),
    ('Middleware', 'CORS for cross-origin requests'),
    ('Deployment', 'localhost:3000 (development)'),
    ('API', 'RESTful API with JSON responses')
]

for i, (component, tech) in enumerate(tech_data, 1):
    row_cells = tech_table.rows[i].cells
    row_cells[0].text = component
    row_cells[1].text = tech

doc.add_paragraph()

# Dependencies Section
doc.add_heading('Project Dependencies', level=1)

doc.add_heading('Production Dependencies:', level=2)
deps = [
    'express (^4.18.2) - Web application framework',
    'cors (^2.8.5) - Cross-Origin Resource Sharing middleware'
]
for dep in deps:
    doc.add_paragraph(dep, style='List Bullet')

doc.add_paragraph()

doc.add_heading('Development Dependencies:', level=2)
dev_deps = [
    'nodemon (^3.0.1) - Automatic server restart on file changes'
]
for dev_dep in dev_deps:
    doc.add_paragraph(dev_dep, style='List Bullet')

doc.add_paragraph()

# Installation Instructions
doc.add_heading('Installation Instructions', level=1)

steps = [
    'Navigate to project directory: cd workspace/app',
    'Install dependencies: npm install',
    'Start the server: npm start',
    'Open browser and navigate to: http://localhost:3000',
    'Fill in the startup idea form with valid data (minimum character requirements)',
    'Click "Analyze Idea" button',
    'Review comprehensive AI analysis results',
    'Click "Clear Results" to analyze another idea'
]

for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph()

# Data Validation Requirements
doc.add_heading('Data Validation Requirements', level=1)

validation_reqs = [
    'Startup Title: Minimum 3 characters, non-empty string',
    'Idea Description: Minimum 10 characters, meaningful content',
    'Target Audience: Minimum 3 characters, non-empty string',
    'All fields are required and validated server-side'
]

for validation in validation_reqs:
    doc.add_paragraph(validation, style='List Bullet')

doc.add_paragraph()

# User Experience Requirements
doc.add_heading('User Experience Requirements', level=1)

ux_reqs = [
    'Responsive design that works on desktop, tablet, and mobile devices',
    'Clear visual feedback during analysis processing (loading spinner)',
    'Informative error messages for invalid inputs',
    'Success message upon successful analysis',
    'Easy-to-read results displayed in organized sections',
    'One-click clear button to reset form and start over',
    'Disabled submit button during processing to prevent duplicate submissions'
]

for ux in ux_reqs:
    doc.add_paragraph(ux, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer = doc.add_paragraph('Document Generated: AI Startup Idea Validator Project Documentation')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_format = footer.runs[0]
footer_format.font.size = Pt(10)
footer_format.font.italic = True
footer_format.font.color.rgb = RGBColor(128, 128, 128)

# Save the document
doc.save('AI_Startup_Idea_Validator_Requirements.docx')
print("✅ Document created successfully: AI_Startup_Idea_Validator_Requirements.docx")
